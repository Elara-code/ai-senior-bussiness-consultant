from io import BytesIO

from docx import Document
from pypdf import PdfReader

from consultant.adapters.export.docx import DocxExporter
from consultant.adapters.export.markdown import MarkdownExporter
from consultant.adapters.export.pdf import PdfExporter
from consultant.ports.exporter import ExportDocument


def test_exporters_include_draft_governance_and_readable_content() -> None:
    source = ExportDocument("Proposal", "Approved scope", False, ["[1] Source page 2"])
    markdown = MarkdownExporter().export(source).content.decode()
    docx = "\n".join(
        p.text for p in Document(BytesIO(DocxExporter().export(source).content)).paragraphs
    )
    pdf = "\n".join(
        page.extract_text() or ""
        for page in PdfReader(BytesIO(PdfExporter().export(source).content)).pages
    )
    assert "未经审批" in markdown
    assert "未经审批" in docx
    assert "DRAFT - NOT APPROVED" in pdf
    assert "Approved scope" in pdf
