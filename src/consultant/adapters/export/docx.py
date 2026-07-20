from io import BytesIO

from docx import Document

from consultant.ports.exporter import ExportDocument, ExportedFile


class DocxExporter:
    def export(self, document: ExportDocument) -> ExportedFile:
        output = BytesIO()
        doc = Document()
        doc.add_heading(document.title, level=1)
        if not document.approved:
            doc.add_paragraph("草稿：未经审批，不得对外使用")
        for line in document.body_markdown.splitlines():
            if line.strip():
                doc.add_paragraph(line.lstrip("# "))
        if document.citation_lines:
            doc.add_heading("引用", level=2)
            for citation in document.citation_lines:
                doc.add_paragraph(citation)
        doc.save(output)
        return ExportedFile(
            output.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "docx",
        )
