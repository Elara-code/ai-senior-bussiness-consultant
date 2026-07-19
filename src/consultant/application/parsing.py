from dataclasses import dataclass
from io import BytesIO

from docx import Document
from pypdf import PdfReader

from consultant.domain.common import ValidationFailure


@dataclass(frozen=True, slots=True)
class ParsedBlock:
    text: str
    section: str | None
    page: int | None


def parse_document(*, payload: bytes, content_type: str) -> list[ParsedBlock]:
    if content_type in {"text/plain", "text/markdown"}:
        return _parse_text(payload, markdown=content_type == "text/markdown")
    if content_type == "application/pdf":
        return _parse_pdf(payload)
    if content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _parse_docx(payload)
    raise ValidationFailure(f"Unsupported content type: {content_type}")


def _parse_text(payload: bytes, *, markdown: bool) -> list[ParsedBlock]:
    try:
        text = payload.decode("utf-8")
    except UnicodeDecodeError as error:
        raise ValidationFailure("Text document must be UTF-8 encoded") from error
    blocks: list[ParsedBlock] = []
    section: str | None = None
    for paragraph in _paragraphs(text):
        if markdown and paragraph.lstrip().startswith("#"):
            section = paragraph.lstrip("# ").strip()
            continue
        blocks.append(ParsedBlock(text=paragraph, section=section, page=None))
    return blocks


def _parse_pdf(payload: bytes) -> list[ParsedBlock]:
    try:
        reader = PdfReader(BytesIO(payload))
        return [
            ParsedBlock(text=paragraph, section=None, page=page_number)
            for page_number, page in enumerate(reader.pages, start=1)
            for paragraph in _paragraphs(page.extract_text() or "")
        ]
    except Exception as error:
        raise ValidationFailure("PDF could not be parsed") from error


def _parse_docx(payload: bytes) -> list[ParsedBlock]:
    try:
        document = Document(BytesIO(payload))
    except Exception as error:
        raise ValidationFailure("DOCX could not be parsed") from error
    blocks: list[ParsedBlock] = []
    section: str | None = None
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            section = text
            continue
        blocks.append(ParsedBlock(text=text, section=section, page=None))
    for table in document.tables:
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        text = "\n".join(row for row in rows if row.strip(" |"))
        if text:
            blocks.append(ParsedBlock(text=text, section=section, page=None))
    return blocks


def _paragraphs(text: str) -> list[str]:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    return [part.strip() for part in normalized.split("\n\n") if part.strip()]
